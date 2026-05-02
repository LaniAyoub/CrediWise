"""Generate CrediWise Chapter 2 Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Page margins ───────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ─── Helpers ────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"),   "single")
        tag.set(qn("w:sz"),    "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color", "BFBFBF"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def heading(text, level):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 0:
        run.font.size  = Pt(16)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0x1F, 0x50, 0x8A)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size  = Pt(14)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0x1F, 0x50, 0x8A)
    elif level == 2:
        run.font.size  = Pt(12)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    else:
        run.font.size  = Pt(11)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_after  = Pt(3)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    return p

def sub_bullet(text):
    p = doc.add_paragraph(style="List Bullet 2")
    p.paragraph_format.left_indent = Cm(2)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text).font.size = Pt(11)
    return p

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size  = Pt(10)
    run.font.bold  = True
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    p.paragraph_format.space_after = Pt(8)
    return p

def figure_placeholder(fig_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ {fig_text} ]")
    run.font.size  = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_table(headers, rows, col_widths=None, header_color="1F508A"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, header_color)
        for run in cell.paragraphs[0].runs:
            run.font.bold  = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size  = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg  = "F2F7FC" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            set_cell_bg(cell, bg)
            set_cell_border(cell, color="BFBFBF")
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def page_break():
    doc.add_page_break()

def space():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)


# ════════════════════════════════════════════════════════════════════════════
#  COVER / CHAPTER TITLE
# ════════════════════════════════════════════════════════════════════════════
heading("Chapitre 2", 0)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Analyse et Spécification des Besoins")
r.font.size = Pt(18)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x50, 0x8A)
doc.add_paragraph()

# ── Table of contents stub ──────────────────────────────────────────────────
heading("Plan", 1)
toc_items = [
    ("2.1", "Spécification des acteurs", "3"),
    ("2.2", "Spécification des besoins", "3"),
    ("2.3", "Backlog de Produit", "5"),
    ("2.4", "Planification des sprints", "8"),
    ("2.5", "Analyse globale du projet", "9"),
    ("2.6", "Architecture du système", "9"),
    ("2.7", "Environnement de développement", "11"),
    ("2.8", "Technologies utilisées", "13"),
]
for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.add_run(f"{num}  {title}").font.size = Pt(11)
    p.add_run(f" {'.' * 60} {page}").font.size = Pt(11)
page_break()

# ════════════════════════════════════════════════════════════════════════════
#  INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════
heading("Introduction", 1)
body(
    "Dans ce chapitre, nous allons identifier toutes les caractéristiques de notre projet CrediWise "
    "en déterminant les besoins fonctionnels et en établissant la liste des exigences non fonctionnelles. "
    "Les acteurs seront définis, et tous les besoins seront clarifiés, puis représentés à travers un "
    "diagramme de cas d'utilisation global. Enfin, nous présenterons la mise en place des différents "
    "sprints pour organiser le développement du projet."
)

# ════════════════════════════════════════════════════════════════════════════
#  2.1  ACTEURS
# ════════════════════════════════════════════════════════════════════════════
heading("2.1  Spécification des acteurs", 1)
heading("2.1.1  Identification des acteurs", 2)
body(
    "L'étude d'une application commence par l'identification de ses différents intervenants. Un acteur "
    "est une personne ou un système qui interagit avec l'application en échangeant des informations. "
    "Dans le cadre du projet CrediWise, plusieurs acteurs sont identifiés selon leurs rôles au sein "
    "de l'établissement bancaire :"
)

actors = [
    ("Super Administrateur (SUPER_ADMIN)",
     "Dispose d'un accès complet à l'ensemble du système. Responsable de la gestion globale de la "
     "plateforme : création et administration des comptes gestionnaires, gestion des agences, "
     "attribution des rôles, et supervision de toutes les entités du système. Il peut également "
     "valider ou rejeter des demandes de crédit."),
    ("Utilisateur Technique (TECH_USER)",
     "Chargé des opérations d'administration technique. Il gère les comptes gestionnaires et les "
     "agences, mais n'a pas accès aux données clients ni aux demandes de crédit."),
    ("Chargé de Relation Client (CRO – Client Relationship Officer)",
     "Responsable de la création et de la mise à jour des fiches clients, ainsi que de la soumission "
     "des demandes de crédit. Il consulte l'état d'avancement des dossiers qui lui sont assignés."),
    ("Agent Front Office",
     "Assure l'accueil et la saisie des dossiers clients au guichet. Il peut créer des fiches clients "
     "et soumettre des demandes de crédit en leur nom."),
    ("Décideur d'Agence (BRANCH_DM)",
     "Responsable de la validation ou du rejet des demandes de crédit soumises au niveau de l'agence, "
     "après examen des dossiers."),
    ("Décideur au Siège (HEAD_OFFICE_DM)",
     "Intervient pour la validation des dossiers de crédit nécessitant une décision au niveau central, "
     "au-delà du seuil de compétence des agences."),
    ("Analyste Risque (RISK_ANALYST)",
     "Spécialiste chargé de l'évaluation du risque associé aux demandes de crédit. Il peut examiner "
     "et rendre un avis sur les dossiers soumis."),
    ("Auditeur (READ_ONLY)",
     "Dispose d'un accès en lecture seule à l'ensemble des données du système, dans le cadre des "
     "missions de contrôle et d'audit interne."),
]
for name, desc in actors:
    bullet(f" {desc}", bold_prefix=f"{name} : ")

# ════════════════════════════════════════════════════════════════════════════
#  2.2  BESOINS
# ════════════════════════════════════════════════════════════════════════════
heading("2.2  Spécification des besoins", 1)
body(
    "Dans cette section, nous analyserons les différents besoins fonctionnels et non fonctionnels, "
    "une étape clé sur laquelle repose la structure de notre projet."
)

heading("2.2.1  Besoins fonctionnels", 2)

func_needs = [
    ("Authentification et sécurité", [
        "Le système est sécurisé avec une authentification par identifiant et mot de passe pour limiter "
        "l'accès aux utilisateurs autorisés selon leurs rôles.",
        "Un mécanisme de déconnexion sécurisée est prévu pour garantir le contrôle permanent des accès.",
        "Les droits d'accès sont gérés de manière granulaire selon le rôle de chaque utilisateur "
        "(contrôle d'accès basé sur les rôles – RBAC).",
    ]),
    ("Gestion des utilisateurs et des agences", [
        "Le Super Administrateur et l'Utilisateur Technique peuvent créer, modifier, activer ou "
        "désactiver des comptes gestionnaires.",
        "Le système permet de gérer les agences bancaires (création, modification) et d'y rattacher "
        "les gestionnaires.",
        "L'attribution des rôles et des permissions est réalisée par le Super Administrateur.",
    ]),
    ("Gestion des clients", [
        "L'agent Front Office et le CRO peuvent créer des fiches clients pour les personnes physiques "
        "et les personnes morales.",
        "Les informations client incluent les données d'identité, de contact, financières, ainsi que "
        "le type de compte, le segment et le secteur d'activité.",
        "Le système permet de rechercher un client par numéro d'identité nationale ou par numéro de "
        "téléphone.",
        "Les données client sont capturées sous forme de snapshot immuable au moment de la soumission "
        "d'une demande.",
    ]),
    ("Gestion des demandes de crédit", [
        "Le CRO et l'agent Front Office peuvent créer des demandes de crédit en état BROUILLON (DRAFT), "
        "en renseignant l'objet du crédit, le montant, la durée, les produits sélectionnés ainsi que "
        "les garanties et garants.",
        "Une demande peut être soumise, ce qui la verrouille pour modification et déclenche la "
        "génération automatique d'un rapport PDF.",
        "Les décideurs (BRANCH_DM, HEAD_OFFICE_DM) et l'analyste risque peuvent valider ou rejeter "
        "les demandes soumises.",
        "Le cycle de vie d'une demande suit les états : BROUILLON → SOUMISE → VALIDÉE / REJETÉE.",
        "Le système offre une pagination et un filtrage des demandes par statut, agence ou identifiant client.",
    ]),
    ("Audit et traçabilité", [
        "Le système enregistre les événements d'authentification (connexion, déconnexion) avec des "
        "métadonnées contextuelles : adresse IP, navigateur, type de terminal, rôle utilisateur.",
        "Les opérations sur les clients (création, mise à jour, suppression, recherche) et sur les "
        "demandes (soumission, validation, changement de statut) sont tracées dans un journal d'événements.",
        "Les utilisateurs en lecture seule (auditeurs) peuvent consulter l'intégralité des données "
        "et des journaux.",
    ]),
    ("Observabilité et supervision", [
        "Le système expose des métriques de performance via Prometheus, consultables sur des tableaux "
        "de bord SigNoz.",
        "Le traçage distribué entre les microservices est assuré par OpenTelemetry, permettant le "
        "suivi des requêtes de bout en bout.",
        "Les journaux structurés au format JSON sont corrélés avec les identifiants de trace pour "
        "faciliter le diagnostic.",
    ]),
]
for title, items in func_needs:
    bullet("", bold_prefix=f"{title} :")
    for item in items:
        sub_bullet(item)

heading("2.2.2  Besoins non fonctionnels", 2)
nfr = [
    ("Performance",
     "L'application est construite sur Quarkus, un framework Java optimisé pour les démarrages rapides "
     "et la faible empreinte mémoire, garantissant des temps de réponse courts même sous forte charge."),
    ("Scalabilité",
     "L'architecture microservices permet de faire évoluer indépendamment chaque service selon les "
     "besoins, sans impacter l'ensemble du système."),
    ("Sécurité",
     "L'application assure la sécurité et la confidentialité des données en utilisant JWT pour "
     "l'authentification, BCrypt pour le hachage des mots de passe, et un contrôle d'accès basé "
     "sur les rôles (RBAC) pour la gestion des autorisations."),
    ("Maintenabilité",
     "L'architecture modulaire en microservices, combinée à Flyway pour les migrations de base de "
     "données et à des tests automatisés, facilite la mise à jour et l'évolution de chaque composant "
     "de manière indépendante."),
    ("Disponibilité",
     "Le déploiement via Docker et Docker Compose assure la portabilité et la reproductibilité de "
     "l'environnement, facilitant les déploiements et réduisant les risques d'indisponibilité."),
    ("Utilisabilité",
     "L'application offre une interface utilisateur intuitive et réactive développée en React, "
     "adaptée aux différents profils d'utilisateurs bancaires, avec un support multilingue via i18next."),
    ("Traçabilité",
     "Toutes les opérations sensibles sont journalisées avec des métadonnées riches, permettant de "
     "répondre aux exigences d'audit interne et de conformité réglementaire."),
]
for name, desc in nfr:
    bullet(f" {desc}", bold_prefix=f"{name} : ")

# ════════════════════════════════════════════════════════════════════════════
#  2.3  BACKLOG
# ════════════════════════════════════════════════════════════════════════════
page_break()
heading("2.3  Backlog de Produit", 1)
body(
    "Après avoir défini les besoins fonctionnels de notre solution, le backlog de produit est "
    "présenté dans le tableau 2.1 ci-après :"
)
caption("Tableau 2.1 : Backlog du produit")

backlog_headers = ["ID", "Fonctionnalité", "User Story", "Complexité", "Priorité"]
backlog_rows = [
    ["1", "Authentification\net sécurité",
     "En tant qu'utilisateur, je veux pouvoir me connecter à l'application avec mes identifiants "
     "afin d'accéder aux fonctionnalités correspondant à mon rôle.",
     "8", "1"],
    ["", "",
     "En tant qu'utilisateur, je veux pouvoir me déconnecter de l'application pour sécuriser mon accès.",
     "5", "1"],
    ["", "",
     "En tant que Super Administrateur, je veux pouvoir attribuer des rôles aux utilisateurs afin "
     "de contrôler leurs permissions sur la plateforme.",
     "8", "1"],
    ["2", "Gestion des\nutilisateurs\net des agences",
     "En tant que Super Administrateur, je veux pouvoir créer, modifier et désactiver des comptes "
     "gestionnaires afin de contrôler les accès à la plateforme.",
     "8", "1"],
    ["", "",
     "En tant que Super Administrateur, je veux pouvoir créer et gérer des agences bancaires et "
     "y rattacher des gestionnaires.",
     "8", "1"],
    ["", "",
     "En tant qu'Utilisateur Technique, je veux pouvoir administrer les comptes gestionnaires et "
     "les agences sans accéder aux données clients.",
     "8", "1"],
    ["3", "Gestion\ndes clients",
     "En tant qu'agent Front Office, je veux pouvoir créer une fiche client (personne physique ou "
     "morale) avec ses informations d'identité, de contact et financières.",
     "8", "1"],
    ["", "",
     "En tant que CRO, je veux pouvoir rechercher un client par numéro d'identité nationale ou "
     "par téléphone afin d'accéder rapidement à son dossier.",
     "5", "1"],
    ["", "",
     "En tant que CRO, je veux pouvoir mettre à jour les informations d'un client existant afin "
     "de maintenir son dossier à jour.",
     "8", "1"],
    ["4", "Gestion des\ndemandes\nde crédit",
     "En tant que CRO, je veux pouvoir créer une demande de crédit en brouillon pour un client "
     "en renseignant l'objet, le montant, la durée et les produits souhaités.",
     "8", "1"],
    ["", "",
     "En tant que CRO, je veux pouvoir ajouter des garanties et des garants à une demande de crédit "
     "afin de renforcer le dossier.",
     "8", "1"],
    ["", "",
     "En tant qu'agent Front Office, je veux pouvoir soumettre une demande de crédit complète afin "
     "qu'elle soit transmise aux décideurs pour examen.",
     "8", "1"],
    ["", "",
     "En tant que décideur, je veux pouvoir consulter la liste des demandes soumises, avec filtrage "
     "par statut et par agence, afin de prioriser mon traitement.",
     "8", "1"],
    ["", "",
     "En tant que décideur, je veux pouvoir valider ou rejeter une demande de crédit soumise après "
     "examen du dossier.",
     "8", "1"],
    ["", "",
     "En tant qu'utilisateur, je veux pouvoir exporter une demande de crédit soumise en format PDF "
     "pour en faire un document formel.",
     "8", "2"],
    ["5", "Audit et\ntraçabilité",
     "En tant qu'auditeur, je veux pouvoir consulter les journaux d'événements d'authentification "
     "et d'opérations afin de vérifier la conformité des accès.",
     "8", "2"],
    ["", "",
     "En tant que Super Administrateur, je veux que toutes les opérations sensibles soient tracées "
     "avec des métadonnées contextuelles (IP, rôle, horodatage).",
     "8", "1"],
    ["6", "Observabilité",
     "En tant qu'administrateur technique, je veux pouvoir consulter les métriques de performance "
     "et les traces distribuées via des tableaux de bord SigNoz afin de superviser la santé du système.",
     "13", "2"],
]
add_table(backlog_headers, backlog_rows, col_widths=[1.2, 3.5, 8.0, 2.0, 1.8])

# ════════════════════════════════════════════════════════════════════════════
#  2.4  SPRINTS
# ════════════════════════════════════════════════════════════════════════════
heading("2.4  Planification des sprints", 1)
body(
    "Le système sera développé en plusieurs versions successives conformément à la méthodologie Scrum. "
    "Chaque version résulte d'une série de sprints, dont l'objectif est de produire progressivement "
    "un ensemble de fonctionnalités qui, une fois combinées, fourniront un produit final offrant une "
    "valeur optimale aux utilisateurs."
)
caption("Tableau 2.2 : Planification des sprints")
sprint_headers = ["Sprint", "Libellé", "Début prévu", "Fin prévu"]
sprint_rows = [
    ["1", "Conception générale et mise en place de l'architecture microservices", "15-02-2024", "02-03-2024"],
    ["",  "Authentification, gestion des rôles et sécurité JWT",                 "03-03-2024", "17-03-2024"],
    ["2", "Gestion des utilisateurs (gestionnaires) et des agences",             "18-03-2024", "30-03-2024"],
    ["",  "Gestion des clients (personnes physiques et morales)",                 "01-04-2024", "21-04-2024"],
    ["3", "Gestion des demandes de crédit (création, soumission, cycle de vie)", "22-04-2024", "22-05-2024"],
    ["",  "Génération PDF et workflow de validation/rejet",                       "23-05-2024", "10-06-2024"],
    ["4", "Audit, traçabilité et observabilité (OpenTelemetry, SigNoz, Prometheus)", "11-06-2024", "30-06-2024"],
]
add_table(sprint_headers, sprint_rows, col_widths=[1.5, 8.5, 2.8, 2.7])

# ════════════════════════════════════════════════════════════════════════════
#  2.5  ANALYSE GLOBALE
# ════════════════════════════════════════════════════════════════════════════
page_break()
heading("2.5  Analyse globale du projet", 1)
body(
    "Une fois la planification du projet établie, nous procédons à la modélisation du projet à travers "
    "des diagrammes qui illustrent les différentes fonctionnalités de l'application."
)

heading("2.5.1  Diagramme de cas d'utilisation général", 2)
body(
    "Les cas d'utilisation jouent un rôle clé dans la représentation des besoins de tout projet. "
    "Dès la phase de spécification, ils servent de référence pour guider le développement d'un outil "
    "qui répond aux attentes des utilisateurs. Dans cette section, nous présentons le diagramme des "
    "cas d'utilisation spécifiques à notre projet CrediWise, illustrant les interactions entre les "
    "différents acteurs et les fonctionnalités du système."
)
figure_placeholder("Figure 2.1 – Insérer ici le diagramme de cas d'utilisation général")
caption("Figure 2.1 : Diagramme de cas d'utilisation général")

heading("2.5.2  Diagramme de classes d'analyse", 2)
body(
    "La figure ci-dessous montre le diagramme de classes d'analyse du système CrediWise. "
    "Il représente les entités principales, leurs attributs, leurs méthodes et les relations "
    "entre les différents composants du domaine métier."
)
figure_placeholder("Figure 2.2 – Insérer ici le diagramme de classes (généré depuis PlantUML)")
caption("Figure 2.2 : Diagramme des classes d'analyse")

# ════════════════════════════════════════════════════════════════════════════
#  2.6  ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════════
heading("2.6  Architecture du système", 1)

heading("2.6.1  Architecture physique", 2)
body(
    "L'architecture physique met en évidence les différents composants matériels et logiciels "
    "supportant l'application. CrediWise repose sur une architecture à trois niveaux basée sur "
    "des microservices conteneurisés."
)
figure_placeholder("Figure 2.3 – Insérer ici le diagramme d'architecture physique")
caption("Figure 2.3 : Architecture Physique")
body("Notre architecture s'appuie sur les composants suivants :")

phys = [
    ("Serveur Frontend",
     "Ce serveur héberge l'application React servie par Vite. Il traite les interactions des "
     "utilisateurs via le protocole HTTP/HTTPS et communique avec les services backend via des "
     "appels API REST."),
    ("Serveurs Backend (Microservices)",
     "Quatre services indépendants s'exécutent sur des ports distincts, chacun exposant une API "
     "REST documentée via Swagger UI : Service Gestionnaire (port 8090), Service Client (port 8082), "
     "Service Nouvelle Demande (port 8083), Service Analyse (port 8084)."),
    ("Serveurs de Bases de Données",
     "Chaque microservice dispose de sa propre instance PostgreSQL (patron Database per Service), "
     "garantissant l'isolation des données : gestionnaire_db (5432), client_db (5433), "
     "nouvelle_demande_db (5434), analyse_db (5435)."),
    ("Plateforme d'Observabilité (SigNoz)",
     "Ce composant centralise la collecte des métriques, des traces distribuées et des journaux "
     "structurés émis par l'ensemble des microservices via le protocole OTLP."),
]
for name, desc in phys:
    bullet(f" {desc}", bold_prefix=f"{name} : ")

heading("2.6.2  Architecture logique", 2)
body(
    "L'architecture logique décrit les composants abstraits, leurs rôles et les interactions entre "
    "eux. Dans le but d'optimiser les performances, la maintenabilité et la réutilisabilité de "
    "l'application, nous avons adopté une architecture en microservices organisée en couches, "
    "comme illustré dans la figure 2.4."
)
figure_placeholder("Figure 2.4 – Insérer ici le diagramme d'architecture logique")
caption("Figure 2.4 : Architecture logique du système")

logic_layers = [
    ("Couche de Présentation",
     "Accessible via un navigateur web, cette couche regroupe les composants React de l'interface "
     "utilisateur. Elle gère la navigation, l'affichage et les interactions, avec un contrôle d'accès "
     "par rôle (RoleGuard) et un support multilingue via i18next."),
    ("Couche API (REST)",
     "Chaque microservice expose des endpoints RESTful documentés via OpenAPI/Swagger. Cette couche "
     "gère la réception des requêtes HTTP, l'authentification JWT, la validation des entrées et la "
     "sérialisation des réponses JSON."),
    ("Couche Métier",
     "Gère la logique applicative propre à chaque domaine : règles de validation, gestion des états "
     "et transitions du cycle de vie des demandes, et orchestration des opérations métier."),
    ("Couche Persistance",
     "S'occupe de l'accès aux bases de données PostgreSQL via Hibernate ORM avec Panache, et assure "
     "la gestion des migrations de schéma avec Flyway. Les échanges inter-services critiques passent "
     "par gRPC."),
]
for name, desc in logic_layers:
    bullet(f" {desc}", bold_prefix=f"{name} : ")

# ════════════════════════════════════════════════════════════════════════════
#  2.7  ENVIRONNEMENT
# ════════════════════════════════════════════════════════════════════════════
page_break()
heading("2.7  Environnement de développement", 1)
heading("2.7.1  Environnement logiciel", 2)
caption("Tableau 2.3 : Environnement logiciel")

env_headers = ["Outil", "Description"]
env_rows = [
    ["Quarkus",
     "Framework Java open source ultraléger et optimisé pour les microservices et les environnements "
     "conteneurisés. Il offre un démarrage ultra-rapide et une faible consommation mémoire, "
     "particulièrement adapté au développement backend des services CrediWise."],
    ["React",
     "Bibliothèque JavaScript open source développée par Meta, utilisée pour construire l'interface "
     "utilisateur de CrediWise. Couplée à TypeScript et Vite, elle permet de développer une SPA "
     "(Single Page Application) réactive et maintenable."],
    ["PostgreSQL",
     "Système de gestion de bases de données relationnelles open source, reconnu pour sa robustesse "
     "et ses performances. Chaque microservice CrediWise dispose de sa propre instance PostgreSQL, "
     "garantissant l'isolation des données."],
    ["Docker / Docker Compose",
     "Plateforme de conteneurisation qui simplifie le déploiement et l'orchestration des microservices "
     "et de leurs dépendances dans des environnements reproductibles."],
    ["Git",
     "Système de contrôle de version décentralisé qui permet à l'équipe de collaborer efficacement "
     "en gérant les branches et les modifications du code source."],
    ["IntelliJ IDEA",
     "Environnement de développement intégré (IDE) polyvalent, particulièrement adapté au "
     "développement Java/Quarkus, offrant des outils intelligents pour coder plus rapidement."],
    ["Visual Studio Code",
     "Éditeur de code flexible et extensible conçu par Microsoft, utilisé pour le développement "
     "frontend React/TypeScript. Il offre la complétion de code, le débogage intégré et l'intégration Git."],
    ["Postman",
     "Client d'API qui simplifie la création, le test et la documentation des APIs REST exposées "
     "par les microservices CrediWise."],
    ["pgAdmin",
     "Interface web d'administration des bases de données PostgreSQL, permettant de visualiser et "
     "d'administrer les instances de bases de données des microservices."],
    ["SigNoz",
     "Plateforme open source d'observabilité qui centralise les métriques Prometheus, les traces "
     "distribuées OpenTelemetry et les journaux structurés de l'ensemble des microservices CrediWise."],
    ["StarUML",
     "Outil de modélisation avancé supportant UML et BPMN, utilisé pour concevoir les diagrammes "
     "de cas d'utilisation, de classes et d'architecture du projet."],
    ["Overleaf",
     "Plateforme collaborative en ligne de rédaction LaTeX, utilisée pour la rédaction du présent "
     "rapport de projet de fin d'études."],
]
add_table(env_headers, env_rows, col_widths=[4.0, 12.5])

# ════════════════════════════════════════════════════════════════════════════
#  2.8  TECHNOLOGIES
# ════════════════════════════════════════════════════════════════════════════
heading("2.8  Technologies utilisées", 1)

techs = [
    ("Quarkus",
     "Quarkus est le framework Java utilisé pour concevoir la partie backend du système CrediWise. "
     "Il offre une solution robuste et performante pour le développement de microservices. Dans notre "
     "projet, Quarkus est combiné avec Hibernate ORM + Panache pour la gestion simplifiée des entités, "
     "Maven pour la gestion des dépendances, Flyway pour les migrations de base de données, et "
     "SmallRye JWT pour la sécurité par tokens."),
    ("React",
     "React est la bibliothèque JavaScript utilisée pour développer la partie frontend de CrediWise. "
     "Elle permet de construire une application web dynamique et interactive sous forme de SPA. Dans "
     "notre projet, React est intégré avec TypeScript pour la rigueur des types, Tailwind CSS pour le "
     "design adaptatif, React Hook Form et Zod pour la gestion des formulaires, Axios pour les appels "
     "API, React Router pour la navigation, et i18next pour le support multilingue."),
    ("gRPC",
     "gRPC est utilisé pour les communications inter-services dans l'architecture microservices de "
     "CrediWise. Il assure des échanges rapides, typés et fiables entre les services backend, notamment "
     "pour la récupération des snapshots clients lors de la création des demandes de crédit."),
    ("OpenTelemetry & SigNoz",
     "OpenTelemetry est le standard ouvert d'instrumentation utilisé dans CrediWise pour collecter "
     "les traces distribuées, les métriques et les journaux de l'ensemble des microservices. SigNoz, "
     "la plateforme d'observabilité, reçoit ces données via le protocole OTLP et les présente sous "
     "forme de tableaux de bord interactifs, permettant la supervision en temps réel de la plateforme."),
]
for name, desc in techs:
    bullet(f" {desc}", bold_prefix=f"{name} : ")
    space()

# ════════════════════════════════════════════════════════════════════════════
#  CONCLUSION
# ════════════════════════════════════════════════════════════════════════════
heading("Conclusion", 1)
body(
    "Dans ce chapitre, nous avons exposé la solution proposée CrediWise en soulignant ses principales "
    "caractéristiques. Les acteurs du système bancaire ont été identifiés selon leurs rôles et niveaux "
    "d'accès, et une analyse complète des besoins fonctionnels et non fonctionnels a été effectuée. "
    "Le backlog de produit et la planification des sprints ont permis d'organiser le développement "
    "de manière itérative et incrémentale, conformément à la méthodologie Scrum. L'architecture "
    "microservices en quatre services indépendants a été présentée, ainsi que l'environnement technique "
    "retenu. Le prochain chapitre abordera la phase de conception détaillée et de réalisation des "
    "premiers sprints."
)

# ════════════════════════════════════════════════════════════════════════════
#  SAVE
# ════════════════════════════════════════════════════════════════════════════
out = r"C:\Users\ranim\OneDrive\Desktop\CrediWise\Chapitre2_CrediWise.docx"
doc.save(out)
print(f"Document saved: {out}")
